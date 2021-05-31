import json
import os
import re
from copy import deepcopy

import numpy as np
from bs4 import NavigableString
from bs4 import Tag
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import Table as DocTable
from treelib import Tree

from src.IO.fileInteraction.FileIO import FileIO
from src.app import gol
from src.tableExtract.TableItem import TableItem
from src.tools.algorithm.exceptionCatch import except_output


def _clearNameOrRel(string: str) -> str:
    """
    清理姓名和关系名,删除符号和括号
    :param string:待处理字符串
    :return:处理完毕的字符串
    """
    if len(string) == 0 or string.isspace():
        return ''
    string = re.sub(u"\(.?\)|\\（.*?）|\\{.*?}|\\[.*?]|\\【.*?】|\\<.*?\\>", "", string)  # 去除括号
    punctuation = "[\s+\.\!\/_,$%^*(+\"\']+|[+——！，。？?、~@#￥%……&*（）]+"
    string = re.sub(punctuation, "", string)  # 去除人名和关系名中的  符号
    return string


def _append(aList: list, a: list, b: str, c: list):
    """
    将a,b,c添加到aList中
    :param aList: 待添加列表
    :param a:[人名：url]
    :param b:关系
    :param c:[人名:url]
    :return:无
    """
    a[0] = _clearNameOrRel(a[0])
    b = _clearNameOrRel(b)
    c[0] = _clearNameOrRel(c[0])
    if len(b) > 7:  # 关系名不能大于七个字符
        return
    if len(a[0]) == 0 or str(a[0]).isspace() or len(b) == 0 or b.isspace() or len(c[0]) == 0 or str(c[0]).isspace():
        return
    if not str(a[0]).isspace() and not str(c[0]).isspace():
        aList.append([a, b, c])


def _notNullAppend(aList: list, a: str, b: str, c: str, isName=False):
    """
    非空添加三元组到列表中，将[a,b,c]添加到列表，如果a，b,c中任意一个为空，则不添加,若长度太长，也不会添加
    :param aList:待添加的列表
    :param a:主体str
    :param b:关系str
    :param c:客体str
    :param isName:第三个属性是否为人名
    :return:无
    """
    if len(a) == 0 or a.isspace() or len(b) == 0 or b.isspace() or len(c) == 0 or c.isspace():
        return
    if len(a) > 7 or len(b) > 7:  # 默认人物名，关系名不超过7个字符
        return
    myList = [a, b, c]
    for i in range(len(myList)):
        myList[i] = re.sub(u"\(.?\)|\\（.*?）|\\{.*?}|\\[.*?]|\\【.*?】||\\<.*?\\>", "", myList[i])  # 去除括号
    punctuation = "[\s+\.\!\/_,$%^*(+\"\']+|[+——！，。？?、~@#￥%……&*（）]+"
    myList[0] = re.sub(punctuation, "", myList[0])  # 去除人名和关系名中的  符号
    myList[1] = re.sub(punctuation, "", myList[1])
    if isName:
        myList[2] = re.sub(punctuation, "", myList[2])
        if len(myList[2]) > 7:
            return
    aList.append([myList[0], myList[1], myList[2]])


class Table:
    """
    表格类
    """

    def __init__(self, rowNumber: int = 0, colNumber: int = 0, name: str = "未命名表格",
                 table=None, unfoldDirection=None):
        """
        初始化函数
        :param rowNumber: 表格的行数
        :param colNumber: 表格的列数
        :param name: 表格的名称
        :param table: 表格的单元格数组，这是一个二维数组
        :param unfoldDirection: 表格的展开方向
        """
        self.rowNumber = rowNumber  # 表格的行数
        self.colNumber = colNumber  # 表格的列数
        if table is None:  # 表格所在的二维数组
            self.cell = [[TableItem(content=0, rowLoc=j, rowspan=1, colLoc=i, colspan=1)
                          for i in range(self.colNumber)]
                         for j in range(self.rowNumber)]
        else:
            self.cell = table
        self.name = name  # 表格的名称
        self.prefix = None  # 表格的前驱
        self.unfoldDirection = unfoldDirection  # 表格的展开方向
        self.__isCorrect = True  # 当前表格是否正确，行列单元数相同
        self.__isNormal = True  # 当前表格是否正常，行列数均大于等于2
        self.propertyList = []  # 属性单元列表
        self.propertyNameList = []  # 属性名列表
        self.propertyLineNum = 1  # 属性行数
        self.tableType = None  # 表格类型
        self.centerWord = None  # 中心词汇,例如人物表，中心词汇就是人名，如“李渊”
        self.hrefMap = {}  # 超链接映射

        self.getAbsolutePosition()  # 获取表格单元的绝对位置
        self.initialNormal()  # 判断表格是否正常
        self.initialTableItemsType()  # 初始化表格单元的类型

    def extendTable(self):
        """
        将当前表格扩展为规范表格
        :return: 扩展后的表格
        """
        # 行扩展
        for rows in self.cell:
            before = 0
            for item in rows:
                if item.rowspan > 1:
                    rowspan = item.rowspan
                    item.rowspan = 1
                    for row in range(item.absoluteRow + 1, item.absoluteRow + rowspan):
                        newItem = deepcopy(item)
                        newItem.rowLoc = row
                        newItem.absoluteRow = row
                        self.cell[row].insert(before, newItem)
                before += 1
        # 列扩展
        for rows in self.cell:
            for item in rows:
                if item.colspan > 1:
                    colspan = item.colspan
                    item.colspan = 1
                    for col in range(item.absoluteCol + 1, item.absoluteCol + colspan):
                        newItem2 = deepcopy(item)
                        newItem2.colLoc = col
                        newItem2.absoluteCol = col
                        self.cell[item.absoluteRow].insert(item.absoluteCol, newItem2)
        self.initialNormal()
        self.initialCorrect()
        return self

    def isCorrect(self):
        """
        判断当前表格是否正确，即行列单元数相同
        :return:
        """
        return self.__isCorrect

    def isNormal(self):
        """
        判断当前表格是否正常，即行列数均大于等于2
        :return:
        """
        return self.__isNormal

    def deleteOneRow(self, index: int):
        """
        删除指定行
        :param index:要删除的索引号，例如Index=0代表第1行
        :return:
        """
        if self.__isCorrect and self.__isNormal:
            if index < 0 or index >= self.rowNumber:
                raise Exception(f"要删除的行<{index}>超出行数范围<0,{self.rowNumber - 1}>")
            del self.cell[index]
            self.rowNumber -= 1
            self.getAbsolutePosition()
            self.initialPropertyList()
        else:
            raise Exception("当前表格未规整，无法删除行")

    def deleteOneCol(self, index: int):
        """
        删除指定列
        :param index: 要删除的索引号，例如Index=0代表第1列
        :return:无
        """
        if self.__isCorrect and self.__isNormal:
            if index < 0 or index >= self.colNumber:
                raise Exception(f"要删除的列<{index}>超出列数范围<0,{self.colNumber - 1}>")
            for i in range(self.rowNumber):
                del self.cell[i][index]
            self.getAbsolutePosition()
            self.colNumber -= 1
            self.initialPropertyList()
        else:
            raise Exception("当前表格未规整，无法删除列")

    def flip(self):
        """
        翻转表格方向,并返回一个新的矩阵
        :return:返回翻转方向后的矩阵
        """
        newTable = Table(rowNumber=self.colNumber, colNumber=self.rowNumber, name=self.name)
        for i in range(self.rowNumber):
            for j in range(self.colNumber):
                newTable.cell[j][i] = deepcopy(self.cell[i][j])
        if self.unfoldDirection == "ROW":
            newTable.unfoldDirection = "COL"
        if self.unfoldDirection == "COL":
            newTable.unfoldDirection = "ROW"

        newTable.prefix = self.prefix  # 表格的前驱
        newTable.propertyList = self.propertyList  # 属性单元列表
        newTable.propertyNameList = self.propertyNameList  # 属性名列表
        newTable.propertyLineNum = self.propertyLineNum  # 属性行数
        newTable.tableType = self.tableType  # 表格类型
        newTable.centerWord = self.centerWord  # 中心词汇,例如人物表，中心词汇就是人名，如“李渊”
        newTable.hrefMap = self.hrefMap  # 超链接映射
        newTable.initialNormal()  # 判断表格是否正常
        newTable.initialCorrect()
        return newTable

    def changeToStr(self):
        """
        将表格内的数据形式全部转化为字符串
        :return:转化后的表格
        """
        for i in range(self.rowNumber):
            for j in range(self.colNumber):
                self.cell[i][j].content = str(self.cell[i][j].content)
        return self

    def getTableItemLengthCharacter(self):
        """
        计算矩阵的几何特征，返回行方差均值和列方差均值，方差越小，则按照该方式展开的可能性越大
        :return: 方差均值和列方差均值
        """
        data = np.zeros((self.rowNumber, self.colNumber), dtype=int)
        for i in range(self.rowNumber):
            for j in range(self.colNumber):
                data[i, j] = len(str(self.cell[i][j].content))
        colVarianceMean = np.mean(np.std(data, axis=0))  # 列方差均值
        rowVarianceMean = np.mean(np.std(data, axis=1))  # 行方差均值
        sumNumber = rowVarianceMean + colVarianceMean
        if sumNumber == 0:
            return rowVarianceMean, colVarianceMean
        return rowVarianceMean / sumNumber, colVarianceMean / sumNumber

    def getTableItemTypeCharacter(self):
        """
        计算矩阵的类型特征，返回行方差均值和列方差均值，方差越小，则按照该方式展开的可能性越大
        :return: 方差均值和列方差均值
        """
        _typeTree = TypeTree()
        return _typeTree.getTypeCharacter(self)

    def getTableItemWordTypeCharacter(self):
        """
        获得行列的单词类型差异
        :return:
        """
        self.initialTableItemWordType()
        for row in self.cell:
            for col in row:
                print(col.wordType, end=" ")
            print()
        data = np.zeros((self.rowNumber, self.colNumber), dtype=int)
        for i in range(self.rowNumber):
            for j in range(self.colNumber):
                data[i, j] = self.cell[i][j].wordType

        colVarianceMean = np.mean(np.std(data, axis=0))  # 列方差均值
        rowVarianceMean = np.mean(np.std(data, axis=1))  # 行方差均值
        sumNumber = rowVarianceMean + colVarianceMean
        if sumNumber == 0:
            return rowVarianceMean, colVarianceMean
        return rowVarianceMean / sumNumber, colVarianceMean / sumNumber

    def getRowAt(self, row: int):
        """
        获取表格第row行的数据列表,如果获取不到则抛出异常
        :param row: 行数，从0开头
        :return: 第row行对应的数据列表
        """
        if self.__isNormal and self.__isCorrect:
            if 0 <= row < self.rowNumber:
                return self.cell[row]
            else:
                raise Exception(f"row={row},此时超出表格索引范围")
        else:
            raise Exception(f"当前表格不正常，无法获取第{row}行的数据列表")

    def getColAt(self, col: int):
        """
        获取表格第col列的数据列表,如果获取不到则抛出异常
        :param col: 列数，从0开头
        :return: 第col列对应的数据列表
        """
        if self.__isNormal and self.__isCorrect:
            if 0 <= col < self.colNumber:
                res = []
                for row in range(self.rowNumber):
                    res.append(self.cell[row][col])
                return res
            else:
                raise Exception(f"col={col},此时超出表格索引范围")
        else:
            raise Exception(f"当前表格不正常，无法获取第{col}列的数据列表")

    def getUnfoldDirection(self) -> str:
        """
        返回表格的展开方向,只能判断为横向展开或者纵向展开
        :return: "ROW"表示横向展开，"COL"表示纵向展开
        """
        if self.unfoldDirection:
            return self.unfoldDirection

        # 标签识别
        rowRes = [item.tagName == 'th' for item in self.getRowAt(0)]
        if rowRes[0] and len(set(rowRes)) == 1:
            self.unfoldDirection = "ROW"
            return self.unfoldDirection
        colRes = [item.tagName == 'th' for item in self.getColAt(0)]
        if colRes[0] and len(set(colRes)) == 1:
            self.unfoldDirection = "COL"
            return self.unfoldDirection

        # 经验规则
        if self.rowNumber >= 3 and self.colNumber >= self.rowNumber * 3:
            self.unfoldDirection = 'ROW'
        if self.colNumber >= 3 and self.rowNumber >= self.colNumber * 3:
            self.unfoldDirection = 'COL'

        # 长度和类型判断法
        rowVarianceMean, colVarianceMean = self.getTableItemLengthCharacter()
        rowTypeCharacter, colTypeCharacter = self.getTableItemTypeCharacter()
        W1 = 0.5
        W2 = 0.5
        Row = W1 * rowVarianceMean + W2 * rowTypeCharacter
        Col = W1 * colVarianceMean + W2 * colTypeCharacter
        if Row < Col:
            direction = "ROW"
        elif Row == Col:
            # 词性和判断法
            rowWordTypeVarianceMean, colWordTypeVarianceMean = self.getTableItemWordTypeCharacter()
            if rowWordTypeVarianceMean < colWordTypeVarianceMean:
                direction = "ROW"
            elif rowWordTypeVarianceMean > colWordTypeVarianceMean:
                direction = "COL"
            else:
                direction = "ROW"  # 如果无法判断，则判断为横向
        else:
            direction = "COL"
        self.unfoldDirection = direction
        return self.unfoldDirection

    def getAbsolutePosition(self):
        """
        获得表格中每个项目所在的绝对位置，其中行绝对位置为self.absoluteRow,列绝对位置为self.absoluteCol
        :return:无
        """
        positionList = []
        for i in range(len(self.cell)):
            colIndex = 0
            before = 0  # 记录从这一行开始，到现在，之前有几个元素进入队列
            for j in range(len(self.cell[i])):
                data = self.cell[i][j]
                colStart = 0
                for position in positionList:
                    colStart += position[1]
                data.absoluteCol = colStart + j - before
                data.absoluteRow = i
                if data.rowspan > 1 or data.colspan > 1:
                    positionList.append([data.rowspan, data.colspan])
                    before += 1
                colIndex += 1

            for x in reversed(range(len(positionList))):
                if positionList[x][0] > 1:
                    positionList[x][0] -= 1
                else:
                    positionList.pop(x)

    def getPropertyList(self, isPropertyName=False) -> list:
        """
        获取属性所在的列表
        :isPropertyName:是否只返回属性名
        :return:属性单元格列表
        """
        if not isPropertyName:
            if self.propertyList:
                return self.propertyList
        else:
            if self.propertyNameList:
                return self.propertyNameList
        self.initialPropertyList()
        if not isPropertyName:
            return self.propertyList
        else:
            self.propertyNameList = [str(item.content) for item in self.propertyList]
            return self.propertyNameList

    def getHrefMap(self) -> dict:
        """
        获取整张表格的href映射表
        :return:整张表格的href映射表
        """
        if len(self.hrefMap) == 0:
            for row in self.cell:
                for col in row:
                    for key in col.href.keys():
                        if key not in self.hrefMap:
                            self.hrefMap[key] = col.href[key]
        else:
            return self.hrefMap

    def getTableType(self):
        """
        识别表格类型
        :return:表格的类型
        """
        if self.tableType:
            return self.tableType
        else:
            if self.__isPersonInfoTable():
                self.tableType = "个人信息表"
            elif self.__isPropertyRelationShipTable():
                self.tableType = "属性关系表"
            elif self.__isTitleRelationShipTable():
                self.tableType = "标题关系表"
            elif self.__isEntityRelationshipTable():
                self.tableType = "实体关系表"
            else:
                self.tableType = "其他表"
        return self.tableType

    def getPersonColList(self, deleteCol=False, removeHeader=False, getName=False) -> list:
        """
        获取人名列表
        :param deleteCol:是否删除人名的这一列
        :param removeHeader:是否去除表头,一般是属性栏
        :param getName: 是否获取人名
        :return:人名的那一列
        """

        def __clearPersonNameList(personNameList: list):
            """
            将人名变成清晰干净的名字
            :param personNameList:
            :return:
            """
            punctuation = "[\s+\.\!\/_,$%^*(+\"\']+|[+——！，。？?、~@#￥%……&*（）]+"
            for i in range(len(personNameList)):
                personNameList[i] = re.sub(u"\(.?\)|\\（.*?）|\\{.*?}|\\[.*?]|\\【.*?】||\\<.*?\\>", "",
                                           personNameList[i])  # 去除括号
                personNameList[i] = str(personNameList[i]).split("/")[0]
                personNameList[i] = re.sub(punctuation, "", personNameList[i])
            return personNameList

        personList = []
        personNameIndex = self.__getPersonNameIndex()
        if personNameIndex != -1:
            personList = [person for person in self.getColAt(personNameIndex)]  # 获得人名所在的表格列
        if len(personList) == 0:
            return personList
        if removeHeader:
            propertyLineNum = self.discriminatePropertyLineNum(self.getUnfoldDirection())
            personList.pop(propertyLineNum - 1)
        if getName:
            personList = [str(person.content) for person in personList]
            personList = __clearPersonNameList(personList)  # 清理人名
        if deleteCol:
            self.deleteOneCol(personNameIndex)
        return personList

    def __tagDiscriminatePropertyLineNum(self, direction: str):
        """
        根据标签判断表格的属性行数，该方法执行前必须先判断表格的展开方向
        :param direction: 表格的展开方向
        :return:
        """
        res = 0
        if direction == "ROW":
            for i in range(self.rowNumber):
                for j in range(self.colNumber):
                    item = self.cell[i][j]
                    if item.tagName != "th":
                        return res
                res += 1
            return res
        elif direction == "COL":
            for j in range(self.colNumber):
                for i in range(self.rowNumber):
                    item = self.cell[i][j]
                    if item.tagName != "th":
                        return res
                res += 1
            return res
        else:
            raise Exception(f"不存在这种表格展开方向<{direction}>")

    def __typeDiscriminatePropertyLineNum(self, direction: str) -> int:
        """
        根据类型判断属性行列数
        :param direction: 展开方向，目前有"ROW"，即行展开，和"COL"，即列展开
        :return: 属性行列数 n，若无法判别，则返回 0
        """
        characterTypeSet = {"字符类型", "中文", "英文", "大写", "小写", "大小写混合"}
        res = 0
        if direction == "ROW":
            for i in range(self.rowNumber):
                for j in range(self.colNumber):
                    item = self.cell[i][j]
                    if item.type_ not in characterTypeSet:
                        return res
                res += 1
            if res == self.rowNumber:  # 如果遍历了所有行
                res = 0
        elif direction == "COL":
            for i in range(self.colNumber):
                for j in range(self.rowNumber):
                    item = self.cell[j][i]
                    if item.type_ not in characterTypeSet:
                        return res
                res += 1
            if res == self.colNumber:  # 如果遍历了所有列
                res = 0
        else:
            raise Exception(f"不存在这种表格展开方向<{direction}>")
        return res

    def discriminatePropertyLineNum(self, direction: str):
        """
        判断表格的属性行数，该方法执行前必须先判断表格的展开方向
        :param direction: 表格的展开方向
        :return:
        """
        if self.propertyLineNum:
            return self.propertyLineNum
        res = self.__tagDiscriminatePropertyLineNum(direction)
        if res == 0 or res > 2:
            res = self.__typeDiscriminatePropertyLineNum(direction)
            if res == 0:
                res = 1
        self.propertyLineNum = res
        return self.propertyLineNum

    def initialTableItemsType(self):
        """
        初始化表格每一个单元的类型，如“你好”就是中文，“123”就是数字>1，“hello”就是英文
        :return:无
        """
        for row in self.cell:
            for item in row:
                item.getTableItemType()

    def initialTableItemWordType(self):
        """
        获得单词类型，例如"水果"就是名词，“跑步”就是动词，如果是句子就会划分为多个词
        :return:无
        """
        for row in self.cell:
            for item in row:
                item.getTableItemWordType()

    def initialCorrect(self) -> bool:
        """
        判断表格是否正确，正确表格的行与列单位数都非常规整
        :return:表格正确则返回True，表格错误则返回False
        """
        colLenList = []
        for rows in self.cell:
            colLen = 0
            for col in rows:
                colLen += col.colspan
            colLenList.append(colLen)
        self.__isCorrect = (len(set(colLenList)) == 1)
        return self.__isCorrect

    def initialNormal(self) -> bool:
        """
        判断是否是一个正常的表格，正常表格必须行列数都大于2
        :return:正常表格则返回True，否则返回False
        """
        if self.rowNumber >= 2 and self.colNumber >= 2:
            self.__isNormal = True
        else:
            self.__isNormal = False
        return self.__isNormal

    def initialPropertyList(self):
        """
        初始化表格的属性列表
        :return: 无
        """
        direction = self.getUnfoldDirection()
        propertyLineNum = self.discriminatePropertyLineNum(direction)
        if direction == "ROW":
            self.propertyList = self.getRowAt(propertyLineNum - 1)
        elif direction == "COL":
            self.propertyList = self.getColAt(propertyLineNum - 1)
        else:
            raise Exception(f"不存在该表格展开方向<{direction}>")
        self.propertyNameList = [str(p.content) for p in self.propertyList]

    def writeTable2Doc(self, filepath: str):
        """
        将表格写入到指定路径的doc文档中
        :param filepath: 指定的文件路径
        :return:无
        """
        if os.path.exists(filepath):
            doc = Document(filepath)
        else:
            doc = Document()
        # 设置字体样式
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # # ------添加文档标题-------
        # paragraph = doc.add_paragraph()
        # run = paragraph.add_run(self.name)
        # font = run.font
        # # 设置字体大小
        # font.size = Pt(12)
        # # 设置水平居中
        # paragraph_format = paragraph.paragraph_format
        # paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        addTable = doc.add_table(rows=self.rowNumber + 1, cols=self.colNumber, style="Table Grid")
        # 添加表格标题
        addTable.cell(0, 0).merge(addTable.cell(0, self.colNumber - 1))
        addTable.cell(0, 0).text = "未命名表格" if self.name is None or self.name == "None" else str(self.name)
        addTable.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 第一行表格水平居中
        self.getAbsolutePosition()
        for rowData in self.cell:
            for item in rowData:
                addTable.cell(item.absoluteRow + 1, item.absoluteCol).merge(
                    addTable.cell(item.absoluteRow + item.rowspan,
                                  item.absoluteCol + item.colspan - 1))
                if item.content is None:
                    addTable.cell(item.absoluteRow + 1, item.absoluteCol).text = "未命名表格"
                else:
                    addTable.cell(item.absoluteRow + 1, item.absoluteCol).text = item.content
        paragraph = doc.add_paragraph()
        doc.save(filepath)

    def writeTable2Pkl(self, filepath: str, mode="wb+"):
        """
        当前表格写入pkl文件
        :param filepath: 文件路径
        :param mode:存储模式
        :return:
        """
        FileIO.writePkl(filepath=filepath, data=self, mode=mode)

    @except_output()
    def __table2DictList(self, filtration=False, deletePersonName=False) -> list:
        """
        表格转化为字典列表,默认为横向展开
        :return:无
        """
        if self.__isNormal and self.__isCorrect:
            if deletePersonName:
                personNameIndex = self.__getPersonNameIndex()
                if personNameIndex != -1:
                    self.deleteOneCol(personNameIndex)
            jsonStrList = []
            direction = self.getUnfoldDirection()
            lineNum = self.discriminatePropertyLineNum(direction)
            if lineNum >= 1:
                heads = [str(head.content) for head in self.getRowAt(lineNum - 1)]
                for i in range(lineNum, self.rowNumber):
                    jsonStr = {}
                    for j in range(self.colNumber):
                        item = self.cell[i][j]
                        if filtration:
                            if str(item.content).isspace() or len(str(item.content)) == 0:
                                continue
                        jsonStr[heads[j]] = str(item.content)
                    jsonStrList.append(jsonStr)
                return jsonStrList
        else:
            raise Exception("该表格不规范，无法写成json串")

    def table2Json(self) -> str:
        """
        表格转化为json串
        :return:转化后的json串
        """
        return json.dumps(self.__table2DictList(), ensure_ascii=False)

    def __isPersonInfoTable(self)->bool:
        """
        识别当前表格是否为人物信息表格
        :return:是则返回True，不是则返回False
        """
        if self.getUnfoldDirection() == "ROW":
            if self.rowNumber != 2:
                return False
        elif self.getUnfoldDirection() == "COL":
            if self.colNumber != 2:
                return False
        firstProperty = {"中文名", "本名"}
        if str(self.cell[0][0].content) in firstProperty:  # 判断第一行第一列的属性名
            self.fusionJsonWord(f"{gol.get_value('personTablePath')}\\personInfo.json")  # 融合属性
            return True

        Threshold = 0.5
        personTablePath = gol.get_value('personTablePath')
        filename = "personInfo.json"
        personProperty = FileIO.readJson(os.path.join(personTablePath, filename))
        personPropertySet = set(personProperty)
        tablePropertySet = set(self.getPropertyList(isPropertyName=True))
        if len(tablePropertySet) == 0:
            proportion = 0
        else:
            proportion = len(tablePropertySet.intersection(personPropertySet)) / len(tablePropertySet)
        if proportion >= Threshold:
            return True
        else:
            return False

    def __isPropertyRelationShipTable(self) -> bool:
        """
        判断是否为属性关系表
        :return:是则返回True，不是则返回False
        """
        # 有列为关系
        personTablePath = gol.get_value('personTablePath')
        propertyRelationShipList = FileIO.readJson(f"{personTablePath}\\propertyRelationship.json")
        propertyList = self.getPropertyList(isPropertyName=True)
        for propertyName in propertyList:
            for relationshipName in propertyRelationShipList:
                if relationshipName in propertyName:
                    return True
        # 属性行为关系
        propertyNameList = self.getPropertyList(isPropertyName=True)
        CRList = FileIO.readJson(f"{gol.get_value('personTablePath')}\\captionRelationship.json")
        count = 0
        for propertyName in propertyNameList:
            for CR in CRList:
                if CR in propertyName:
                    count += 1
                    continue
        if count > len(propertyNameList) / 2:
            return True
        return False

    def __isTitleRelationShipTable(self) -> bool:
        """
        判断是否为标题关系表
        :return:是则返回True，否则返回False
        """
        if self.name:
            personTablePath = gol.get_value('personTablePath')
            relationshipList = FileIO.readJson(f"{personTablePath}\\captionRelationship.json")
            for relationship in relationshipList:
                if relationship in self.name:
                    return True
            # reg = re.compile(r".*[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+.*")
            # if re.match(reg, str(self.name)):
            #     return True
        return False

    def __isEntityRelationshipTable(self) -> bool:
        """
        判断是否为实体关系表
        :return:是则返回True，否则返回False
        """
        personTablePath = gol.get_value('personTablePath')
        relationshipList = FileIO.readJson(f"{personTablePath}\\personName.json")
        propertyList = self.getPropertyList(isPropertyName=True)
        for propertyName in propertyList:
            for relationship in relationshipList:
                if relationship in propertyName:
                    return True
        return False

    def __getPropertyRelationshipList(self):
        """
        获取属性关系列表，并且把与人物有关的属性由高到低排序
        :return:属性关系列表
        """
        personTablePath = gol.get_value('personTablePath')
        propertyRelationshipList = FileIO.readJson(f"{personTablePath}\\propertyRelationship.json")
        propertyList = self.getPropertyList(isPropertyName=True)
        indexAndNameList = []
        for propertyName in propertyList:
            indexAndNameList.extend(
                [(index, relationshipName) for index, relationshipName
                 in list(enumerate(propertyRelationshipList))
                 if relationshipName in propertyName])
        sortIndexList = sorted(indexAndNameList, key=lambda indexAndNum: indexAndNum[0])  # 根据序号排序
        sortIndexList = [indexAndName[1] for indexAndName in sortIndexList]
        return sortIndexList

    def fusionJsonWord(self, jsonFilePath):
        """
        该函数暂时未用到。
        融合个人属性到“个人信息表”列表之中,使得下一次的判断更加精确
        :return:
        """
        personProperty = FileIO.readJson(jsonFilePath)
        personPropertySet = set(personProperty)
        tablePropertySet = set(self.getPropertyList(isPropertyName=True))
        personPropertySet = personPropertySet.union(tablePropertySet)
        FileIO.write2Json(list(personPropertySet), jsonFilePath)

    @except_output("从表格中抽取实体和关系时出错")
    def extractEntityRelationship(self, getEntityTriad=False):
        """
        抽取实体关系
        :return:从当前表格中抽取的实体列表和关系列表
        """
        entity = []
        relationship = []
        typeName = self.getTableType()
        if typeName == "个人信息表" or typeName == "实体关系表":
            entity = self.extractEntity(getEntityTriad)
        elif typeName == "属性关系表":
            relationship = self.extractPropertyRelationship()
            entity = self.extractEntity(getEntityTriad)
        elif typeName == "标题关系表":
            relationship = self.extractCaptionRelationship()
            entity = self.extractEntity(getEntityTriad)
        else:  # 其他表
            # self.writeTable2Doc(f"{gol.get_value('tableDocPath')}\\未抽取三元组的表格.docx")
            pass
        return entity, relationship

    def extractPropertyRelationship(self):
        """
        从当前表格中抽取属性关系
        :return: 属性关系列表
        """
        def listFindPosition(AList: list, waitFind: str):
            for i in range(len(AList)):
                if waitFind in AList[i]:
                    return i
            return -1

        relationship = []
        if not self.prefix:
            return relationship
        propertyNameList = self.getPropertyList(isPropertyName=True)
        if len(propertyNameList) == 0:
            return relationship
        propertyRelationshipList = self.__getPropertyRelationshipList()  # 属性关系列表，例如[关系,辈分]
        if len(propertyRelationshipList) == 0:  # 如果不存在属性关系，则返回空
            return relationship
        if len(propertyRelationshipList) >= 1:  # 如果存在多个属性关系，则删除其余低级的属性关系
            for i in range(1, len(propertyRelationshipList)):
                self.deleteOneCol(listFindPosition(propertyNameList, propertyRelationshipList[i]))
            propertyNameList = self.getPropertyList(isPropertyName=True)  # 删除属性列之后更新一下属性列表
        personNameList = self.getPersonColList()
        personHrefList = self.__getPersonHrefList(personNameList)
        index = listFindPosition(propertyNameList, propertyRelationshipList[0])
        relationshipList = [str(relationship.content) for relationship in self.getColAt(index)]  # 获得关系名列表
        self.deleteOneCol(index)  # 删除关系名列表
        propertyLineNum = self.discriminatePropertyLineNum(self.getUnfoldDirection())
        prefix = [self.prefix, self.hrefMap[self.prefix] if self.prefix in self.hrefMap else '']
        for i in range(propertyLineNum, self.rowNumber):
            # 构建三元组
            if i < len(relationshipList) and i < len(personHrefList):
                _append(relationship, prefix, relationshipList[i], personHrefList[i])
        return relationship

    def extractCaptionRelationship(self):
        """
        从表格中抽取标题关系
        :return:标题关系列表
        """
        relationship = []
        if self.name and self.prefix:
            personNameList = self.getPersonColList(removeHeader=True)
            prefix = [self.prefix, self.hrefMap[self.prefix] if self.prefix in self.hrefMap else '']
            if len(personNameList) == 0:
                propertyNameList = self.getPropertyList(isPropertyName=True)
                CRList = FileIO.readJson(f"{gol.get_value('personTablePath')}\\captionRelationship.json")
                count = 0
                for propertyName in propertyNameList:
                    for CR in CRList:
                        if CR in propertyName:
                            count += 1
                            continue
                if count > len(propertyNameList) / 2:
                    for j in range(self.colNumber):
                        item = self.cell[1][j]
                        if item.href and str(item.content) in item.href:
                            nameAndHref = [str(item.content), item.href[str(item.content)]]
                        else:
                            nameAndHref = [str(item.content), '']
                        _append(relationship, prefix, propertyNameList[j], nameAndHref)
                return relationship
            personHrefList = self.__getPersonHrefList(personNameList)
            for i in range(len(personNameList)):
                # 添加三元组
                _append(relationship, prefix, self.name, personHrefList[i])
        return relationship

    def __extractPropertyRelationship(self):
        """
        抽取属性行为关系的表格
        :return:
        """
        relationship = []
        propertyNameList = self.getPropertyList(isPropertyName=True)
        CRList = FileIO.readJson(f"{gol.get_value('personTablePath')}\\captionRelationship.json")
        count = 0
        for propertyName in propertyNameList:
            for CR in CRList:
                if CR in propertyName:
                    count += 1
                    continue
        if count > len(propertyNameList) / 2:
            for j in range(self.colNumber):
                item = self.cell[1][j]
                if item.href and str(item.content) in item.href:
                    nameAndHref = [str(item.content), item.href[str(item.content)]]
                else:
                    nameAndHref = [str(item.content), '']
                _append(relationship, self.prefix, propertyNameList[j], nameAndHref)
        return relationship

    def extractEntity(self, getEntityTriad=False):
        """
        从表格中抽取实体
        :return:实体列表
        """
        entity = []
        if getEntityTriad:
            personNameList = self.getPersonColList(deleteCol=True, getName=True)  # 获取并删除人名
            if self.colNumber >= 1 and len(personNameList) != 0:
                propertyIndex = self.discriminatePropertyLineNum(self.getUnfoldDirection()) - 1
                propertyNameList = self.getPropertyList(isPropertyName=True)  # 获取除了人名之外的，其余的属性列表
                propertyLine = self.discriminatePropertyLineNum(self.getUnfoldDirection())  # 属性行所占的行数
                for i in range(propertyLine, self.rowNumber):
                    for j in range(propertyIndex, self.colNumber):
                        content = str(self.cell[i][j].content)
                        # 添加三元组
                        _notNullAppend(entity, personNameList[i], propertyNameList[j], content)
        else:
            # if self.colNumber == 1:  # 仅剩一个属性了，必然不存在实体
            #     return entity
            personNameList = self.getPersonColList(getName=True, removeHeader=True)
            if len(personNameList) == 0:
                return entity
            personHrefList = self.__getPersonHrefList(self.getPersonColList(removeHeader=True))
            for i in range(len(personHrefList)):
                personHrefList[i][0] = _clearNameOrRel(personHrefList[i][0])
            dictList = self.__table2DictList(filtration=True, deletePersonName=True)
            if len(personNameList) == len(dictList):
                for i in range(len(personNameList)):
                    if len(personNameList[i]) == 0 or str(personNameList[i]).isspace():  # 姓名为空则跳过
                        continue
                    entity.append([personHrefList[i], dictList[i]])

        return entity

    def __getPersonHrefList(self, personList: list):
        """
        返回人的href链接
        :param personList:代表人物的单元格列表
        :return:人物的href链接链表，每个元素都是一个超链接字典
        """

        personHrefList = []
        for person in personList:
            personName = str(person.content)
            href = person.href
            if personName in href:
                personHrefList.append([personName, href[personName]])
            elif personName in self.hrefMap:
                personHrefList.append([personName, self.hrefMap[personName]])
            else:
                personHrefList.append([personName, ''])
        return personHrefList

    def __getPersonNameIndex(self):
        """
        返回人名所在的列的索引
        :return:人名所在的列的索引
        """

        def _getListIndex(name: str, stringList: list):
            """
            从指定列表中获取包含某个字符串的索引号
            :param name: 指定字符串
            :param stringList: 指定列表
            :return: 索引号，若无，则返回-1
            """
            __index = 0
            for string in stringList:
                if string in name:
                    return __index
                else:
                    __index += 1
            return -1

        personNameIndex = -1
        personTablePath = gol.get_value('personTablePath')
        relationshipList = FileIO.readJson(f"{personTablePath}\\personName.json")
        propertyList = self.getPropertyList(isPropertyName=True)
        index = 0
        for propertyName in propertyList:
            personNameIndex = _getListIndex(propertyName, relationshipList)
            if personNameIndex != -1:
                personNameIndex = index
                break
            index += 1
        return personNameIndex

    def clearTable(self):
        """
        清理表格，去除表格中无意义的序号，去除空行或者空列
        :return:无
        """
        propertyList = self.getPropertyList(isPropertyName=True)
        # 清除带有“序”的属性行
        clearSet = ["序号", "序"]
        indexes = [index for index, propertyName in enumerate(propertyList) if propertyName in clearSet]
        if indexes:
            if self.getUnfoldDirection() == "ROW":
                self.deleteOneCol(indexes[0])
            else:
                self.deleteOneRow(indexes[0])
            self.getAbsolutePosition()
        # 如果第一行内容为空，则删除第一行
        canContinue = True
        for item in self.getRowAt(self.rowNumber - 1):
            if not (len(str(item.content)) == 1 or str(item.content).isspace()):
                canContinue = False
            if not canContinue:
                break
        if canContinue:
            self.deleteOneRow(self.rowNumber - 1)
        # 如果最后一行是参考资料，删除这一行
        canContinue = True
        for item in self.getRowAt(self.rowNumber - 1):
            if "参考资料" not in str(item.content):
                canContinue = False
            if not canContinue:
                break
        if canContinue:
            self.deleteOneRow(self.rowNumber - 1)

        # 将表格中的单纯的符号单元转化为空格
        for i in range(self.rowNumber):
            for j in range(self.colNumber):
                if self.cell[i][j].getTableItemType() == "标点类型":
                    self.cell[i][j].content = ""

    def dump(self):
        string = ""
        for row in self.cell:
            string += ",".join([str(item.content) for item in row]) + "\n"
        return string


class TypeTree:
    """
    类型树类
    """

    def __init__(self):
        """
        初始化类型树
        """
        tree = Tree()
        tree.create_node(tag="类型", identifier="类型")
        tree.create_node(tag="超链接", identifier="超链接", parent="类型")
        tree.create_node(tag="图片", identifier="图片", parent="类型")
        tree.create_node(tag="字符和数字", identifier="字符和数字", parent="类型")
        tree.create_node(tag="其他类型", identifier="其他类型", parent="类型")
        tree.create_node(tag="标点类型", identifier="标点类型", parent="类型")
        tree.create_node(tag="字符类型", identifier="字符类型", parent="字符和数字")
        tree.create_node(tag="数字类型", identifier="数字类型", parent="字符和数字")
        tree.create_node(tag="中文", identifier="中文", parent="字符类型")
        tree.create_node(tag="英文", identifier="英文", parent="字符类型")
        tree.create_node(tag="<=0", identifier="<=0", parent="数字类型")
        tree.create_node(tag="0-1", identifier="0-1", parent="数字类型")
        tree.create_node(tag=">=1", identifier=">=1", parent="数字类型")
        tree.create_node(tag="大写", identifier="大写", parent="英文")
        tree.create_node(tag="小写", identifier="小写", parent="英文")
        tree.create_node(tag="大小写混合", identifier="大小写混合", parent="英文")
        # tree.show()
        self.tree = tree

    def getTypeCharacter(self, table: Table):
        """
        计算表格的行类型特征和列类型特征
        :param table:传入的表格
        :return:行类型特征rowTypeCharacter和列类型特征colTypeCharacter
        """

        # data1 = np.zeros((table.rowNumber, table.colNumber - 1), dtype=int)
        # for i in range(table.rowNumber):
        #     for j in range(table.colNumber - 1):
        #         data1[i][j] = self._VType(table.cell[i][j], table.cell[i][table.colNumber - 1])
        #
        # data2 = np.zeros((table.rowNumber - 1, table.colNumber), dtype=int)
        # for i in range(table.rowNumber - 1):
        #     for j in range(table.colNumber):
        #         data2[i][j] = self._VType(table.cell[i][j], table.cell[table.rowNumber - 1][j])
        # colTypeCharacter = np.mean(data1, axis=0)  # 列方差均值
        # rowTypeCharacter = np.mean(data2, axis=1)  # 行方差均值

        rowTypeCharacter = 0
        colTypeCharacter = 0
        rowTypeCharacterList = []
        colTypeCharacterList = []
        for i in range(table.rowNumber - 1):
            colTypeCharacterList.append(self.VType(table.getRowAt(i), table.getRowAt(table.rowNumber - 1)))
        if colTypeCharacterList:
            colTypeCharacter = np.mean(colTypeCharacterList)

        for j in range(table.colNumber - 1):
            rowTypeCharacterList.append(self.VType(table.getColAt(j), table.getColAt(table.colNumber - 1)))
        if rowTypeCharacterList:
            rowTypeCharacter = np.mean(rowTypeCharacterList)
        sumNumber = rowTypeCharacter + colTypeCharacter
        if sumNumber == 0:
            return rowTypeCharacter, colTypeCharacter
        return rowTypeCharacter / sumNumber, colTypeCharacter / sumNumber

    def _VType(self, item1: TableItem, item2: TableItem) -> int:
        """
        计算两个表格单元之间的类型差异距离
        :param item1: 表格单元1
        :param item2: 表格单元2
        :return: 类型差异距离
        """
        distance = 0
        typeNode1 = item1.type_
        typeNode2 = item2.type_
        if typeNode1 is None or typeNode2 is None:
            raise Exception("当前类型为None，无法计算出类型之间的距离")
        level1 = self.tree.depth(typeNode1)
        level2 = self.tree.depth(typeNode2)
        if typeNode1 == typeNode2:
            return distance
        if level1 > level2:
            while level1 != level2:
                typeNode1 = self.tree.parent(typeNode1).identifier
                distance += 1
                level1 -= 1
        elif level2 > level1:
            while level1 != level2:
                typeNode2 = self.tree.parent(typeNode2).identifier
                distance += 1
                level2 -= 1
        if level1 == level2:
            while typeNode1 != typeNode2:
                typeNode1 = self.tree.parent(typeNode1).identifier
                typeNode2 = self.tree.parent(typeNode2).identifier
                distance += 2
        return distance

    def VType(self, v1: list, v2: list) -> float:
        """
        计算两个列表之间的类型差异
        :param v1:列表1
        :param v2:列表2
        :return:类型差异值
        """
        res = 0
        len1 = len(v1)
        len2 = len(v2)
        if len1 == 0 or len2 == 0:
            return res
        m = min(len1, len2)
        for i in range(m):
            res += self._VType(v1[i], v2[i])
        return res / m


def changeTig2Table(tag: Tag, caption='未命名表格', prefix=None) -> Table:
    """
    将tag标签转化为table数据结构
    :param prefix: 前缀
    :param caption:标题
    :param tag: 输入的标签
    :return: 返回表格
    """

    def changeTag2TableItem(tag: Tag, rowIndex: int, colIndex: int) -> TableItem:
        """
        把标签转化为单元格
        :param tag: 带转化标签
        :param rowIndex: 单元格的行索引
        :param colIndex: 单元格的列索引
        :return: 单元格
        """
        rowspan = colspan = 1
        # 获取表格单元中的超链接
        href = {}
        aList = tag.find_all("a")
        for a in aList:
            if a.has_attr("href"):
                href[a.text] = r"https://baike.baidu.com" + a["href"]
        # 获取表格单元中的图片
        imgSrc = []
        imgList = tag.find_all("img")
        for img in imgList:
            if img.has_attr("src"):
                imgSrc.append(img["src"])
        # 获取表格的占据行列数
        if tag.has_attr("rowspan"):
            rowspan = int(tag['rowspan'])
        if tag.has_attr("colspan"):
            colspan = int(tag['colspan'])
        text = re.sub('(\[)\d+(\])', '', tag.text)  # 去除索引注释，例如 [12]
        content = text.replace("\xa0", "")
        tagName = tag.name
        tableItem = TableItem(content, rowIndex, rowspan, colIndex, colspan, href, imgSrc, tagName=tagName)
        return tableItem

    def finalDeal(table: Table, colLenList: list, rowNumber: int):
        """
        最终处理，该步骤将表格重新初始化，并且重新计算绝对位置，判断表格类型
        :param table: 表格名
        :param colLenList: 列长度列表
        :param rowNumber: 行数
        :return:
        """
        table.colNumber = max(colLenList)
        table.rowNumber = rowNumber
        table.getAbsolutePosition()
        table.initialNormal()  # 判断是否正常
        table.initialCorrect()  # 判断是否正确
        table.initialTableItemsType()  # 初始化表格单元的类型

    table = Table()
    table.cell = []
    colLenList = []
    colIndex = rowIndex = 0
    table.name = str(caption)  # 命名
    table.prefix = prefix
    thead = tag.find("thead")
    tbody = tag.find("tbody")
    if thead and tbody:
        rowIndex = 0
        for row in thead.children:
            colIndex = 0
            colSize = 0
            innerList = []
            for colData in row.children:
                tableItem = changeTag2TableItem(colData, rowIndex, colIndex)
                colIndex += 1
                innerList.append(tableItem)
                colSize += tableItem.colspan
            colLenList.append(colSize)
            table.cell.append(innerList)
            rowIndex += 1
        for row in tbody.children:
            colIndex = 0
            colSize = 0
            innerList = []
            for colData in row.children:
                tableItem = changeTag2TableItem(colData, rowIndex, colIndex)
                colIndex += 1
                innerList.append(tableItem)
                colSize += tableItem.colspan
            rowIndex += 1
            colLenList.append(colSize)
            table.cell.append(innerList)
    else:
        for rowData in tag.children:
            innerList = []
            colSize = 0
            colIndex = 0
            for colData in rowData.children:
                if isinstance(colData, NavigableString):
                    continue
                tableItem = changeTag2TableItem(colData, rowIndex, colIndex)
                colIndex += 1
                colSize += tableItem.colspan
                innerList.append(tableItem)
            colLenList.append(colSize)
            table.cell.append(innerList)
            rowIndex += 1
    finalDeal(table, colLenList, rowIndex)
    # print(table.extendTable().dump())
    return table


def changeWordTable2Table(table: DocTable) -> Table:
    """
    将word表格转化为Table表格
    :param table: word中的表格
    :return: 自定义的Table表
    """
    caption = str(table.rows[0].cells[0].text)
    maxColNum = 0
    rowList = []
    rowNum = len(table.rows)
    for i in range(1, rowNum):
        row = table.rows[i]
        colList = []
        colNum = len(table.rows[i].cells)
        maxColNum = max(maxColNum, colNum)
        for j in range(colNum):
            item = table.rows[i].cells[j]
            newTableItem = TableItem(item.text, i - 1, 1, j, 1)
            colList.append(newTableItem)
        rowList.append(colList)
    newTable = Table(rowNum - 1, maxColNum, caption, rowList)
    return newTable
